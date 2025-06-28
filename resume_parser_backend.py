from flask import Flask, Blueprint, request, jsonify, send_from_directory, current_app, flash, redirect, url_for, session
from flask_cors import CORS
import os
import json
import uuid # To generate unique filenames for temporary files
from dotenv import load_dotenv # Import load_dotenv
from users import UserInDB, get_database_client
from auth import get_password_hash, verify_password, get_user

import re
import base64
from typing import Dict, List, Any
import logging
from datetime import datetime
from io import BytesIO

# Document processing imports
import fitz  # PyMuPDF for PDF processing
from docx import Document
from PIL import Image

# NLP imports
import spacy

# Mistral API client
from mistralai import Mistral

# Set up logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

app = Flask(__name__)
CORS(app)
app.secret_key = os.getenv('FLASK_SECRET_KEY', 'super_secret_key_for_dev') # Add a secret key for sessions

# Load environment variables from .env file
load_dotenv()

# Configuration
UPLOAD_FOLDER = 'uploads'
ALLOWED_EXTENSIONS = {'pdf', 'docx', 'doc', 'png', 'jpg', 'jpeg', 'tiff', 'bmp'}
MAX_CONTENT_LENGTH = 16 * 1024 * 1024  # 16MB max file size

app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['MAX_CONTENT_LENGTH'] = MAX_CONTENT_LENGTH

# Create upload directory if it doesn't exist
os.makedirs(UPLOAD_FOLDER, exist_ok=True)

# Mistral API configuration
MISTRAL_API_KEY = os.getenv('MISTRAL_API_KEY')
if not MISTRAL_API_KEY:
    logger.warning("MISTRAL_API_KEY not found in environment variables")

# Initialize Mistral client
mistral_client = None
if MISTRAL_API_KEY:
    try:
        mistral_client = Mistral(api_key=MISTRAL_API_KEY)
    except Exception as e:
        logger.error(f"Failed to initialize Mistral client: {str(e)}")

# Load spaCy model
try:
    nlp = spacy.load("en_core_web_sm")
except OSError:
    logger.error("spaCy model 'en_core_web_sm' not found. Please install it using: python -m spacy download en_core_web_sm")
    nlp = None

class MistralOCRResumeParser:
    """Enhanced resume parser using Mistral's OCR and layout understanding capabilities"""
    
    def __init__(self):
        self.section_keywords = {
            'contact': ['contact', 'personal information', 'contact information', 'contact details'],
            'summary': ['summary', 'profile', 'objective', 'about', 'overview', 'professional summary'],
            'education': ['education', 'academic', 'qualification', 'degree', 'university', 'college', 'school'],
            'experience': ['experience', 'work', 'employment', 'career', 'professional', 'job', 'work history'],
            'skills': ['skills', 'technical skills', 'competencies', 'abilities', 'technologies', 'core competencies'],
            'projects': ['projects', 'project', 'portfolio', 'work samples', 'key projects'],
            'certifications': ['certifications', 'certificates', 'license', 'credentials', 'professional certifications'],
            'achievements': ['achievements', 'awards', 'honors', 'accomplishments', 'recognition'],
            'languages': ['languages', 'language skills', 'linguistic abilities']
        }
        
        # Enhanced regex patterns
        self.email_pattern = re.compile(r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b')
        self.phone_pattern = re.compile(r'(\+?\d{1,3}[-.\s]?)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}')
        self.linkedin_pattern = re.compile(r'linkedin\.com/in/[\w-]+', re.IGNORECASE)
        self.github_pattern = re.compile(r'github\.com/[\w-]+', re.IGNORECASE)
        
    def allowed_file(self, filename: str) -> bool:
        """Check if file extension is allowed"""
        return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS
    
    def convert_pdf_to_images(self, file_path: str) -> List[Image.Image]:
        """Convert PDF pages to images for OCR processing"""
        images = []
        try:
            doc = fitz.open(file_path)
            for page_num in range(len(doc)):
                page = doc.load_page(page_num)
                pix = page.get_pixmap(matrix=fitz.Matrix(2.0, 2.0))  # 2x zoom for better quality
                img_data = pix.tobytes("png")
                img = Image.open(BytesIO(img_data))
                images.append(img)
            doc.close()
            return images
        except Exception as e:
            logger.error(f"Error converting PDF to images: {str(e)}")
            return []

    def extract_links_from_pdf(self, file_path: str) -> List[str]:
        """Extract hyperlinks from a PDF document using PyMuPDF."""
        links = []
        try:
            doc = fitz.open(file_path)
            for page_num in range(len(doc)):
                page = doc.load_page(page_num)
                page_links = page.get_links()
                for link in page_links:
                    if link['kind'] == fitz.LINK_URI:
                        links.append(link['uri'])
            doc.close()
            return list(set(links)) # Return unique links
        except Exception as e:
            logger.error(f"Error extracting links from PDF: {str(e)}")
            return []
    
    def convert_docx_to_images(self, file_path: str) -> List[Image.Image]:
        """Convert DOCX to images (simplified approach)"""
        # For DOCX, we'll extract text first and create a simple layout
        # In a production environment, you might want to use python-docx2pdf + pdf2image
        try:
            doc = Document(file_path)
            text_content = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)
            
            # Create a simple image representation (placeholder)
            # In practice, you'd want to preserve actual layout
            full_text = '\n'.join(text_content)
            return [self.create_text_image(full_text)]
        except Exception as e:
            logger.error(f"Error converting DOCX to images: {str(e)}")
            return []
    
    def create_text_image(self, text: str) -> Image.Image:
        """Create a simple text image (placeholder for actual DOCX rendering)"""
        from PIL import ImageDraw, ImageFont
        
        # Create a white background image
        img = Image.new('RGB', (800, 1000), color='white')
        draw = ImageDraw.Draw(img)
        
        try:
            font = ImageFont.load_default()
        except:
            font = None
        
        # Draw text
        lines = text[:1000].split('\n')  # Limit text
        y_position = 10
        for line in lines[:30]:  # Limit lines
            draw.text((10, y_position), line, fill='black', font=font)
            y_position += 25
        
        return img
    
    def image_to_base64(self, image: Image.Image) -> str:
        """Convert PIL Image to base64 string"""
        buffered = BytesIO()
        image.save(buffered, format="PNG")
        img_str = base64.b64encode(buffered.getvalue()).decode()
        return img_str
    
    def extract_text_with_mistral_ocr(self, images: List[Image.Image]) -> Dict[str, Any]:
        """Use Mistral's OCR capabilities to extract structured text with layout awareness"""
        if not mistral_client or not images:
            return {"error": "Mistral client not available or no images provided"}
        
        try:
            all_extracted_data = []
            
            for i, image in enumerate(images):
                # Convert image to base64
                image_base64 = self.image_to_base64(image)
                
                # Create the prompt for context-aware OCR
                ocr_prompt = r"""
                You are an expert resume parser with advanced OCR capabilities. Analyze this resume image and extract structured information while preserving layout context.

                Please extract and organize the following information in JSON format:
                {
                    "page_number": <page_number>,
                    "layout_analysis": {
                        "sections_detected": ["list of section headers found"],
                        "layout_type": "single_column|multi_column|complex",
                        "text_regions": ["description of main text regions"]
                    },
                    "contact_information": {
                        "name": "full name",
                        "email": "email address",
                        "phone": "phone number",
                        "linkedin": "linkedin profile",
                        "github": "github profile",
                        "location": "city, state/country",
                        "website": "personal website",
                        "urls": ["list of all detected URLs/hyperlinks"]
                    },
                    "professional_summary": "summary or objective text",
                    "experience": [
                        {
                            "position": "job title",
                            "company": "company name",
                            "duration": "start date - end date",
                            "location": "location",
                            "achievements": ["list of achievements/responsibilities"],
                            "technologies": ["technologies used"]
                        }
                    ],
                    "education": [
                        {
                            "degree": "degree type and major",
                            "institution": "school/university name",
                            "graduation_date": "date",
                            "gpa": "if mentioned",
                            "relevant_coursework": ["if mentioned"]
                        }
                    ],
                    "skills": {
                        "technical_skills": ["programming languages, frameworks, etc."],
                        "soft_skills": ["communication, leadership, etc."],
                        "tools_and_technologies": ["specific tools and platforms"]
                    },
                    "projects": [
                        {
                            "name": "project name",
                            "description": "project description",
                            "technologies": ["technologies used"],
                            "links": ["github, demo links"]
                        }
                    ],
                    "certifications": [
                        {
                            "name": "certification name",
                            "issuer": "issuing organization",
                            "date": "date obtained",
                            "expiry": "expiry date if applicable"
                        }
                    ],
                    "achievements_and_awards": ["list of achievements"],
                    "languages": ["list of languages with proficiency if mentioned"]
                }

                Pay special attention to:
                1. Preserve the hierarchical structure and relationships between elements
                2. Identify section boundaries accurately based on visual layout
                3. Extract dates in a consistent format
                4. Separate different types of skills appropriately
                5. Maintain the context and meaning of bullet points and achievements
                6. Handle multi-column layouts correctly

                Only include sections that are actually present in the resume. If a section is empty or not found, omit it from the JSON.
                """
                
                # Create message with image
                messages = [
                    {
                        "role": "user",
                        "content": [
                            {
                                "type": "text",
                                "text": ocr_prompt
                            },
                            {
                                "type": "image_url",
                                "image_url": f"data:image/png;base64,{image_base64}"
                            }
                        ]
                    }
                ]
                
                # Call Mistral API
                response = mistral_client.chat.complete(
                    model="pixtral-12b-2409",  # Mistral's vision model
                    messages=messages,
                    max_tokens=4000,
                    temperature=0.1
                )
                
                # Parse the response
                response_text = response.choices[0].message.content
                
                # Try to extract JSON from the response
                try:
                    # Find JSON in the response
                    json_start = response_text.find('{')
                    json_end = response_text.rfind('}') + 1
                    if json_start != -1 and json_end != -1:
                        json_str = response_text[json_start:json_end]
                        page_data = json.loads(json_str)
                        page_data['page_number'] = i + 1
                        all_extracted_data.append(page_data)
                    else:
                        # Fallback: include raw text
                        all_extracted_data.append({
                            'page_number': i + 1,
                            'raw_text': response_text,
                            'parsing_error': 'Could not extract structured JSON'
                        })
                except json.JSONDecodeError as e:
                    logger.error(f"JSON parsing error: {str(e)}")
                    all_extracted_data.append({
                        'page_number': i + 1,
                        'raw_text': response_text,
                        'parsing_error': f'JSON decode error: {str(e)}'
                    })
            
            return {
                'pages': all_extracted_data,
                'total_pages': len(images),
                'extraction_method': 'mistral_ocr'
            }
            
        except Exception as e:
            logger.error(f"Mistral OCR extraction error: {str(e)}")
            return {"error": f"Mistral OCR failed: {str(e)}"}
    
    def merge_multi_page_data(self, pages_data: List[Dict[str, Any]]) -> Dict[str, Any]:
        """Merge data from multiple pages into a single resume structure"""
        merged_data = {
            'contact_information': {},
            'professional_summary': '',
            'experience': [],
            'education': [],
            'skills': {'technical_skills': [], 'soft_skills': [], 'tools_and_technologies': []},
            'projects': [],
            'certifications': [],
            'achievements_and_awards': [],
            'languages': [],
            'layout_analysis': {
                'total_pages': len(pages_data),
                'sections_detected': set(),
                'layout_types': []
            }
        }
        
        for page_data in pages_data:
            if 'parsing_error' in page_data:
                continue
                
            # Merge contact information (usually on first page)
            if page_data.get('contact_information') and not merged_data['contact_information']:
                merged_data['contact_information'] = page_data['contact_information']
            
            # Merge professional summary
            if page_data.get('professional_summary') and not merged_data['professional_summary']:
                merged_data['professional_summary'] = page_data['professional_summary']
            
            # Merge lists (extend)
            for list_field in ['experience', 'education', 'projects', 'certifications', 'achievements_and_awards', 'languages']:
                if page_data.get(list_field):
                    if isinstance(page_data[list_field], list):
                        merged_data[list_field].extend(page_data[list_field])
                    else:
                        merged_data[list_field].append(page_data[list_field])
            
            # Merge skills
            if page_data.get('skills'):
                skills = page_data['skills']
                for skill_type in ['technical_skills', 'soft_skills', 'tools_and_technologies']:
                    if skills.get(skill_type):
                        merged_data['skills'][skill_type].extend(skills[skill_type])
            
            # Merge layout analysis
            if page_data.get('layout_analysis'):
                layout = page_data['layout_analysis']
                if layout.get('sections_detected'):
                    merged_data['layout_analysis']['sections_detected'].update(layout['sections_detected'])
                if layout.get('layout_type'):
                    merged_data['layout_analysis']['layout_types'].append(layout['layout_type'])
        
        # Convert sets to lists for JSON serialization
        merged_data['layout_analysis']['sections_detected'] = list(merged_data['layout_analysis']['sections_detected'])
        
        # Remove duplicates from skills
        for skill_type in merged_data['skills']:
            merged_data['skills'][skill_type] = list(set(merged_data['skills'][skill_type]))
        
        return merged_data
    
    def parse_resume(self, file_path: str, filename: str) -> Dict[str, Any]:
        """Main function to parse resume using Mistral OCR"""
        try:
            # Extract file extension
            file_extension = filename.rsplit('.', 1)[1].lower() if '.' in filename else ''
            
            # Convert document to images based on file type
            images = []
            # Removed extracted_urls = []
            if file_extension == 'pdf':
                images = self.convert_pdf_to_images(file_path)
                # Removed extracted_urls = self.extract_links_from_pdf(file_path)
            elif file_extension in ['docx', 'doc']:
                images = self.convert_docx_to_images(file_path)
            elif file_extension in ['png', 'jpg', 'jpeg', 'tiff', 'bmp']:
                images = [Image.open(file_path)]
            
            if not images:
                return {"error": "Could not process the file or extract images"}
            
            # Use Mistral OCR for extraction
            if mistral_client:
                ocr_result = self.extract_text_with_mistral_ocr(images)
                
                if 'error' in ocr_result:
                    return ocr_result
                
                # Merge multi-page data
                if ocr_result.get('pages'):
                    merged_data = self.merge_multi_page_data(ocr_result['pages'])
                    
                    merged_data['extraction_metadata'] = {
                        'total_pages': ocr_result['total_pages'],
                        'extraction_method': 'mistral_ocr',
                        'processing_timestamp': datetime.now().isoformat()
                    }
                    return merged_data
                else:
                    return {"error": "No data extracted from pages"}
            else:
                return {"error": "Mistral OCR service not available"}
            
        except Exception as e:
            logger.error(f"Error parsing resume: {str(e)}")
            return {"error": f"Error parsing resume: {str(e)}"}


    
    def convert_docx_to_images(self, file_path: str) -> List[Image.Image]:
        """Convert DOCX to images (simplified approach)"""
        # For DOCX, we'll extract text first and create a simple layout
        # In a production environment, you might want to use python-docx2pdf + pdf2image
        try:
            doc = Document(file_path)
            text_content = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)
            
            # Create a simple image representation (placeholder)
            # In practice, you'd want to preserve actual layout
            full_text = '\n'.join(text_content)
            return [self.create_text_image(full_text)]
        except Exception as e:
            logger.error(f"Error converting DOCX to images: {str(e)}")
            return []
    
    def create_text_image(self, text: str) -> Image.Image:
        """Create a simple text image (placeholder for actual DOCX rendering)"""
        from PIL import ImageDraw, ImageFont
        
        # Create a white background image
        img = Image.new('RGB', (800, 1000), color='white')
        draw = ImageDraw.Draw(img)
        
        try:
            font = ImageFont.load_default()
        except:
            font = None
        
        # Draw text
        lines = text[:1000].split('\n')  # Limit text
        y_position = 10
        for line in lines[:30]:  # Limit lines
            draw.text((10, y_position), line, fill='black', font=font)
            y_position += 25
        
        return img
    
    def image_to_base64(self, image: Image.Image) -> str:
        """Convert PIL Image to base64 string"""
        buffered = BytesIO()
        image.save(buffered, format="PNG")
        img_str = base64.b64encode(buffered.getvalue()).decode()
        return img_str
    
    def extract_text_with_mistral_ocr(self, images: List[Image.Image]) -> Dict[str, Any]:
        """Use Mistral's OCR capabilities to extract structured text with layout awareness"""
        if not mistral_client or not images:
            return {"error": "Mistral client not available or no images provided"}
        
        try:
            all_extracted_data = []
            
            for i, image in enumerate(images):
                # Convert image to base64
                image_base64 = self.image_to_base64(image)
                
                # Create the prompt for context-aware OCR
                ocr_prompt = r"""
                You are an expert resume parser with advanced OCR capabilities. Analyze this resume image and extract structured information while preserving layout context.

                Please extract and organize the following information in JSON format:
                {
                    "page_number": <page_number>,
                    "layout_analysis": {
                        "sections_detected": ["list of section headers found"],
                        "layout_type": "single_column|multi_column|complex",
                        "text_regions": ["description of main text regions"]
                    },
                    "contact_information": {
                        "name": "full name",
                        "email": "email address",
                        "phone": "phone number",
                        "linkedin": "linkedin profile",
                        "github": "github profile",
                        "location": "city, state/country",
                        "website": "personal website",
                        "urls": ["list of all detected URLs/hyperlinks"]
                    },
                    "professional_summary": "summary or objective text",
                    "experience": [
                        {
                            "position": "job title",
                            "company": "company name",
                            "duration": "start date - end date",
                            "location": "location",
                            "achievements": ["list of achievements/responsibilities"],
                            "technologies": ["technologies used"]
                        }
                    ],
                    "education": [
                        {
                            "degree": "degree type and major",
                            "institution": "school/university name",
                            "graduation_date": "date",
                            "gpa": "if mentioned",
                            "relevant_coursework": ["if mentioned"]
                        }
                    ],
                    "skills": {
                        "technical_skills": ["programming languages, frameworks, etc."],
                        "soft_skills": ["communication, leadership, etc."],
                        "tools_and_technologies": ["specific tools and platforms"]
                    },
                    "projects": [
                        {
                            "name": "project name",
                            "description": "project description",
                            "technologies": ["technologies used"],
                            "links": ["github, demo links"]
                        }
                    ],
                    "certifications": [
                        {
                            "name": "certification name",
                            "issuer": "issuing organization",
                            "date": "date obtained",
                            "expiry": "expiry date if applicable",
                            "link": "link to certification if available"
                        }
                    ],
                    "achievements_and_awards": ["list of achievements"],
                    "languages": ["list of languages with proficiency if mentioned"]
                }

                Pay special attention to:
                1. Preserve the hierarchical structure and relationships between elements
                2. Identify section boundaries accurately based on visual layout
                3. Extract dates in a consistent format
                4. Separate different types of skills appropriately
                5. Maintain the context and meaning of bullet points and achievements
                6. Handle multi-column layouts correctly
                7. **Crucially, identify and extract all hyperlinks/URLs present in the document, especially from contact information, project sections, and certifications.**

                Only include sections that are actually present in the resume. If a section is empty or not found, omit it from the JSON.
                """
                
                # Create message with image
                messages = [
                    {
                        "role": "user",
                        "content": [
                            {
                                "type": "text",
                                "text": ocr_prompt
                            },
                            {
                                "type": "image_url",
                                "image_url": f"data:image/png;base64,{image_base64}"
                            }
                        ]
                    }
                ]
                
                # Call Mistral API
                response = mistral_client.chat.complete(
                    model="pixtral-12b-2409",  # Mistral's vision model
                    messages=messages,
                    max_tokens=4000,
                    temperature=0.1
                )
                
                # Parse the response
                response_text = response.choices[0].message.content
                
                # Try to extract JSON from the response
                try:
                    # Find JSON in the response
                    json_start = response_text.find('{')
                    json_end = response_text.rfind('}') + 1
                    if json_start != -1 and json_end != -1:
                        json_str = response_text[json_start:json_end]
                        page_data = json.loads(json_str)
                        page_data['page_number'] = i + 1
                        all_extracted_data.append(page_data)
                    else:
                        # Fallback: include raw text
                        all_extracted_data.append({
                            'page_number': i + 1,
                            'raw_text': response_text,
                            'parsing_error': 'Could not extract structured JSON'
                        })
                except json.JSONDecodeError as e:
                    logger.error(f"JSON parsing error: {str(e)}")
                    all_extracted_data.append({
                        'page_number': i + 1,
                        'raw_text': response_text,
                        'parsing_error': f'JSON decode error: {str(e)}'
                    })
            
            return {
                'pages': all_extracted_data,
                'total_pages': len(images),
                'extraction_method': 'mistral_ocr'
            }
            
        except Exception as e:
            logger.error(f"Mistral OCR extraction error: {str(e)}")
            return {"error": f"Mistral OCR failed: {str(e)}"}
    
    def merge_multi_page_data(self, pages_data: List[Dict[str, Any]]) -> Dict[str, Any]:
        """Merge data from multiple pages into a single resume structure"""
        merged_data = {
            'contact_information': {},
            'professional_summary': '',
            'experience': [],
            'education': [],
            'skills': {'technical_skills': [], 'soft_skills': [], 'tools_and_technologies': []},
            'projects': [],
            'certifications': [],
            'achievements_and_awards': [],
            'languages': [],
            'layout_analysis': {
                'total_pages': len(pages_data),
                'sections_detected': set(),
                'layout_types': []
            }
        }
        
        for page_data in pages_data:
            if 'parsing_error' in page_data:
                continue
                
            # Merge contact information (usually on first page)
            if page_data.get('contact_information') and not merged_data['contact_information']:
                merged_data['contact_information'] = page_data['contact_information']
            
            # Merge professional summary
            if page_data.get('professional_summary') and not merged_data['professional_summary']:
                merged_data['professional_summary'] = page_data['professional_summary']
            
            # Merge lists (extend)
            for list_field in ['experience', 'education', 'projects', 'certifications', 'achievements_and_awards', 'languages']:
                if page_data.get(list_field):
                    if isinstance(page_data[list_field], list):
                        merged_data[list_field].extend(page_data[list_field])
                    else:
                        merged_data[list_field].append(page_data[list_field])
            
            # Merge skills
            if page_data.get('skills'):
                skills = page_data['skills']
                for skill_type in ['technical_skills', 'soft_skills', 'tools_and_technologies']:
                    if skills.get(skill_type):
                        merged_data['skills'][skill_type].extend(skills[skill_type])
            
            # Merge layout analysis
            if page_data.get('layout_analysis'):
                layout = page_data['layout_analysis']
                if layout.get('sections_detected'):
                    merged_data['layout_analysis']['sections_detected'].update(layout['sections_detected'])
                if layout.get('layout_type'):
                    merged_data['layout_analysis']['layout_types'].append(layout['layout_type'])
        
        # Convert sets to lists for JSON serialization
        merged_data['layout_analysis']['sections_detected'] = list(merged_data['layout_analysis']['sections_detected'])
        
        # Remove duplicates from skills
        for skill_type in merged_data['skills']:
            merged_data['skills'][skill_type] = list(set(merged_data['skills'][skill_type]))
        
        return merged_data
    
    def parse_resume(self, file_path: str, filename: str) -> Dict[str, Any]:
        """Main function to parse resume using Mistral OCR"""
        try:
            # Extract file extension
            file_extension = filename.rsplit('.', 1)[1].lower() if '.' in filename else ''
            
            # Convert document to images based on file type
            images = []
            if file_extension == 'pdf':
                images = self.convert_pdf_to_images(file_path)
            elif file_extension in ['docx', 'doc']:
                images = self.convert_docx_to_images(file_path)
            elif file_extension in ['png', 'jpg', 'jpeg', 'tiff', 'bmp']:
                images = [Image.open(file_path)]
            
            if not images:
                return {"error": "Could not process the file or extract images"}
            
            # Use Mistral OCR for extraction
            if mistral_client:
                ocr_result = self.extract_text_with_mistral_ocr(images)
                
                if 'error' in ocr_result:
                    return ocr_result
                
                # Merge multi-page data
                if ocr_result.get('pages'):
                    merged_data = self.merge_multi_page_data(ocr_result['pages'])
                    merged_data['extraction_metadata'] = {
                        'total_pages': ocr_result['total_pages'],
                        'extraction_method': 'mistral_ocr',
                        'processing_timestamp': datetime.now().isoformat()
                    }
                    return merged_data
                else:
                    return {"error": "No data extracted from pages"}
            else:
                return {"error": "Mistral OCR service not available"}
            
        except Exception as e:
            logger.error(f"Error parsing resume: {str(e)}")
            return {"error": f"Error parsing resume: {str(e)}"}

# Initialize parser
parser = MistralOCRResumeParser()

@app.route('/')
def index():
    """Serve the main HTML page"""
    return send_from_directory('.', 'index.html')

@app.route('/signup', methods=['POST'])
def signup():
    username = request.form.get('username')
    email = request.form.get('email')
    password = request.form.get('password')

    if not username or not email or not password:
        flash('Please fill in all fields.', 'error')
        return redirect(url_for('index'))

    client = get_database_client()
    db = client["ai_resume"]
    users_collection = db["user"]

    if users_collection.find_one({"username": username}):
        flash('Username already exists. Please choose a different one.', 'error')
        client.close()
        return redirect(url_for('index'))
    
    if users_collection.find_one({"email": email}):
        flash('Email already registered. Please use a different email or login.', 'error')
        client.close()
        return redirect(url_for('index'))

    hashed_password = get_password_hash(password)
    user_data = {
        "username": username,
        "email": email,
        "hashed_password": hashed_password,
        "disabled": False
    }
    users_collection.insert_one(user_data)
    client.close()
    flash('Signup successful! Please log in.', 'success')
    return redirect(url_for('index'))

@app.route('/login', methods=['POST'])
def login():
    username = request.form.get('username')
    password = request.form.get('password')

    if not username or not password:
        flash('Please enter both username and password.', 'error')
        return redirect(url_for('index'))

    user = get_user(username)
    if not user or not verify_password(password, user.hashed_password):
        flash('Invalid username or password.', 'error')
        return redirect(url_for('index'))

    flash('Login successful!', 'success')
    session['username'] = user.username # Store username in session
    return redirect(url_for('dashboard'))

@app.route('/dashboard')
def dashboard():
    """Serve the dashboard HTML page"""
    return send_from_directory('.', 'dashboard.html')

@app.route('/resume_parser')
def resume_parser_page():
    """Serve the resume parser HTML page"""
    if 'username' not in session:
        flash('Please log in to access the resume parser.', 'error')
        return redirect(url_for('index'))
    return send_from_directory('.', 'resume_parser.html')

@app.route('/my_info')
def my_info_page():
    """Serve the my info HTML page"""
    # Temporarily bypass authentication for testing
    # if 'username' not in session:
    #     flash('Please log in to view your information.', 'error')
    #     return redirect(url_for('index'))
    return send_from_directory('.', 'my_info.html')

@app.route('/resume_builder')
def resume_builder_page():
    """Serve the resume builder HTML page"""
    if 'username' not in session:
        flash('Please log in to access the resume builder.', 'error')
        return redirect(url_for('index'))
    return send_from_directory('.', 'resume_builder.html')

@app.route('/my_info_data')
def my_info_data():
    """Fetch parsed resume data for the logged-in user"""
    # Temporarily bypass authentication for testing
    # if 'username' not in session:
    #     return jsonify({'error': 'Unauthorized. Please log in.'}), 401
    
    # For testing, use a dummy username if not logged in
    username = session.get('username', 'test_user') 
    client = get_database_client()
    db = client["ai_resume"]
    parsed_resumes_collection = db["parsed_resumes"]
    
    user_resumes = []
    for resume in parsed_resumes_collection.find({"username": username}).sort("timestamp", -1):
        resume['_id'] = str(resume['_id']) # Convert ObjectId to string
        user_resumes.append(resume)
    client.close()
    
    return jsonify({'resumes': user_resumes}), 200

@app.route('/logout')
def logout():
    session.pop('username', None)
    flash('You have been logged out.', 'info')
    return redirect(url_for('index'))

@app.route('/upload', methods=['POST'])
def upload_file():
    """Handle file upload"""
    if 'username' not in session:
        return jsonify({'error': 'Unauthorized. Please log in.'}), 401

    try:
        if 'file' not in request.files:
            return jsonify({'error': 'No file provided'}), 400
        
        file = request.files['file']
        if file.filename == '':
            return jsonify({'error': 'No file selected'}), 400
        
        if file and parser.allowed_file(file.filename):
            # Save file
            filename = f"{datetime.now().strftime('%Y%m%d_%H%M%S')}_{file.filename}"
            file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
            file.save(file_path)
            
            return jsonify({
                'message': 'File uploaded successfully',
                'filename': filename,
                'file_path': file_path
            }), 200
        else:
            return jsonify({'error': 'File type not allowed'}), 400
            
    except Exception as e:
        logger.error(f"Upload error: {str(e)}")
        return jsonify({'error': f'Upload failed: {str(e)}'}), 500

@app.route('/parse', methods=['POST'])
def parse_resume():
    """Parse uploaded resume using Mistral OCR and store in DB"""
    if 'username' not in session:
        return jsonify({'error': 'Unauthorized. Please log in.'}), 401

    try:
        data = request.get_json()
        if not data or 'filename' not in data:
            return jsonify({'error': 'Filename not provided'}), 400
        
        filename = data['filename']
        file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        
        if not os.path.exists(file_path):
            return jsonify({'error': 'File not found'}), 404
        
        # Check if Mistral API key is available
        if not MISTRAL_API_KEY:
            return jsonify({'error': 'Mistral API key not configured'}), 503
        
        # Parse the resume using Mistral OCR
        parsed_data = parser.parse_resume(file_path, filename)
        
        # Clean up uploaded file
        try:
            os.remove(file_path)
        except OSError:
            pass
        
        # Return parsed data without saving to DB here
        return jsonify(parsed_data), 200
        
    except Exception as e:
        logger.error(f"Parse error: {str(e)}")
        return jsonify({'error': f'Parsing failed: {str(e)}'}), 500

@app.route('/save_parsed_resume', methods=['POST'])
def save_parsed_resume():
    """Save manually edited parsed resume data to DB"""
    if 'username' not in session:
        return jsonify({'error': 'Unauthorized. Please log in.'}), 401
    
    username = session['username']
    data = request.get_json()

    if not data:
        return jsonify({'error': 'No data provided for saving.'}), 400

    client = get_database_client()
    db = client["ai_resume"]
    parsed_resumes_collection = db["parsed_resumes"]

    # For simplicity, we'll save it as a new entry.
    # In a real application, you might want to update an existing entry.
    resume_entry = {
        "username": username,
        "filename": "manual_edit_" + datetime.now().strftime('%Y%m%d_%H%M%S'),
        "timestamp": datetime.now(),
        "parsed_data": data
    }
    parsed_resumes_collection.insert_one(resume_entry)
    client.close()
    logger.info(f"Manually edited resume data for {username} saved to DB.")
    
    return jsonify({'message': 'Parsed data saved successfully!'}), 200

@app.route('/delete_resume/<resume_id>', methods=['DELETE'])
def delete_resume(resume_id):
    """Delete a parsed resume entry from the database"""
    if 'username' not in session:
        return jsonify({'error': 'Unauthorized. Please log in.'}), 401
    
    username = session['username']
    client = get_database_client()
    db = client["ai_resume"]
    parsed_resumes_collection = db["parsed_resumes"]

    try:
        from bson.objectid import ObjectId
        # Ensure the resume belongs to the logged-in user before deleting
        result = parsed_resumes_collection.delete_one(
            {"_id": ObjectId(resume_id), "username": username}
        )
        
        if result.deleted_count == 1:
            logger.info(f"Resume {resume_id} deleted by {username}.")
            return jsonify({'message': 'Resume deleted successfully!'}), 200
        else:
            return jsonify({'error': 'Resume not found or not authorized to delete.'}), 404
    except Exception as e:
        logger.error(f"Error deleting resume {resume_id}: {str(e)}")
        return jsonify({'error': f'Failed to delete resume: {str(e)}'}), 500

@app.route('/update_resume/<resume_id>', methods=['PUT'])
def update_resume(resume_id):
    """Update a parsed resume entry in the database"""
    if 'username' not in session:
        return jsonify({'error': 'Unauthorized. Please log in.'}), 401
    
    username = session['username']
    data = request.get_json()

    if not data or 'parsed_data' not in data:
        return jsonify({'error': 'No parsed_data provided for updating.'}), 400

    client = get_database_client()
    db = client["ai_resume"]
    parsed_resumes_collection = db["parsed_resumes"]

    try:
        from bson.objectid import ObjectId
        # Update the document, ensuring it belongs to the logged-in user
        result = parsed_resumes_collection.update_one(
            {"_id": ObjectId(resume_id), "username": username},
            {"$set": {"parsed_data": data['parsed_data'], "last_updated": datetime.now()}}
        )
        
        if result.matched_count == 1:
            if result.modified_count == 1:
                logger.info(f"Resume {resume_id} updated by {username}.")
                return jsonify({'message': 'Resume updated successfully!'}), 200
            else:
                return jsonify({'message': 'No changes detected or saved.'}), 200
        else:
            return jsonify({'error': 'Resume not found or not authorized to update.'}), 404
    except Exception as e:
        logger.error(f"Error updating resume {resume_id}: {str(e)}")
        return jsonify({'error': f'Failed to update resume: {str(e)}'}), 500

@app.route('/parse-preview', methods=['POST'])
def parse_resume_preview():
    """Preview parse results without full processing (for testing)"""
    try:
        data = request.get_json()
        if not data or 'filename' not in data:
            return jsonify({'error': 'Filename not provided'}), 400
        
        filename = data['filename']
        file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        
        if not os.path.exists(file_path):
            return jsonify({'error': 'File not found'}), 404
        
        # Basic file info
        file_stats = os.stat(file_path)
        file_extension = filename.rsplit('.', 1)[1].lower() if '.' in filename else ''
        
        # Try to get basic structure without full OCR processing
        preview_data = {
            'filename': filename,
            'file_size': file_stats.st_size,
            'file_type': file_extension,
            'mistral_ocr_available': mistral_client is not None,
            'processing_ready': MISTRAL_API_KEY is not None,
            'estimated_pages': 1  # Default, would be calculated from actual file
        }
        
        # If PDF, get page count
        if file_extension == 'pdf':
            try:
                doc = fitz.open(file_path)
                preview_data['estimated_pages'] = len(doc)
                doc.close()
            except:
                pass
        
        return jsonify(preview_data), 200
        
    except Exception as e:
        logger.error(f"Preview error: {str(e)}")
        return jsonify({'error': f'Preview failed: {str(e)}'}), 500

@app.route('/templates')
def get_templates():
    """Fetch available resume templates from the database"""
    if 'username' not in session:
        return jsonify({'error': 'Unauthorized. Please log in.'}), 401

    client = get_database_client()
    db = client["ai_resume"]
    templates_collection = db["resume_templates"]

    templates = []
    for template in templates_collection.find({}, {"template_id": 1, "_id": 0}): # Only fetch template_id
        templates.append(template['template_id'])
    client.close()
    return jsonify({'templates': templates}), 200

@app.route('/preview_resume', methods=['POST'])
def preview_resume():
    data = request.json
    template_id = data.get('template_id')
    parsed_data_id = data.get('parsed_data_id')

    if not template_id or not parsed_data_id:
        return jsonify({"error": "Template ID and Parsed Data ID are required."}), 400

    client = get_database_client()
    db = client["ai_resume"]
    templates_collection = db["resume_templates"]
    parsed_resumes_collection = db["parsed_resumes"]

    try:
        from bson.objectid import ObjectId
        # Fetch template content
        template_data = templates_collection.find_one({"template_id": template_id})
        if not template_data:
            client.close()
            return jsonify({"error": "Template not found."}), 404

        html_content = template_data.get("html_content", "")
        css_content = template_data.get("css_content", "")
        js_content = template_data.get("js_content", "")

        # Fetch parsed resume data
        parsed_resume = parsed_resumes_collection.find_one({"_id": ObjectId(parsed_data_id)})
        if not parsed_resume:
            client.close()
            return jsonify({"error": "Parsed resume data not found."}), 404

        parsed_data = parsed_resume.get("parsed_data", {})

        # Construct the full HTML for the iframe
        # Escape </script> tags in js_content to prevent premature termination
        escaped_js_content = js_content.replace("</script>", "<\\/script>")

        iframe_html_content = f"""
            <!DOCTYPE html>
            <html lang="en">
            <head>
                <meta charset="UTF-8">
                <meta name="viewport" content="width=device-width, initial-scale=1.0">
                <title>Resume Preview</title>
                <style>
                    body {{ font-family: sans-serif; margin: 20px; }}
                    .editable-field {{ border: 1px solid #ccc; padding: 5px; margin: 2px 0; width: calc(100% - 12px); }}
                    .editable-list-item {{ margin-bottom: 10px; padding: 10px; border: 1px dashed #eee; }}
                    .add-item-button, .remove-item-button {{
                        background-color: #007bff; color: white; border: none; padding: 5px 10px;
                        border-radius: 3px; cursor: pointer; margin-top: 5px;
                    }}
                    .remove-item-button {{ background-color: #dc3545; margin-left: 5px; }}
                    {css_content}
                </style>
            </head>
            <body>
                {html_content}
                <script type="application/json" id="resumeDataJson">
                    {json.dumps(parsed_data)}
                </script>
                <script>
                    window.resumeData = JSON.parse(document.getElementById('resumeDataJson').textContent);
                    window.parsedDataId = "{parsed_data_id}";

                    // Function to populate placeholders and make them editable
                    // Helper function to get a value from an object based on a path string
                    function getObjectValue(obj, path) {{
                        return path.split('.').reduce((o, i) => (o ? o[i] : undefined), obj);
                    }}

                    // Helper function to set a value in an object based on a path string
                    function setObjectValue(obj, path, value) {{
                        const parts = path.split('.');
                        let current = obj;
                        for (let i = 0; i < parts.length - 1; i++) {{
                            if (!current[parts[i]]) {{
                                current[parts[i]] = {{}};
                            }}
                            current = current[parts[i]];
                        }}
                        current[parts[parts.length - 1]] = value;
                    }}

                    // Function to populate placeholders and make them editable
                    function renderEditableResume() {{
                        // Iterate over all elements with data-path attribute
                        document.querySelectorAll('[data-path]').forEach(element => {{
                            const dataPath = element.dataset.path;
                            let value = getObjectValue(window.resumeData, dataPath);

                            // Handle array values (e.g., skills, achievements)
                            if (Array.isArray(value)) {{
                                value = value.join('\\n'); // Join array items with newline for textarea
                            }} else if (typeof value === 'object' && value !== null) {{
                                // If it's an object, we might need to stringify or handle specifically
                                // For now, we'll just skip or show a placeholder
                                value = JSON.stringify(value); // Or handle more gracefully
                            }} else {{
                                value = value || ''; // Ensure value is a string
                            }}

                            if (element.tagName === 'INPUT' || element.tagName === 'TEXTAREA') {{
                                element.value = escapeHtml(value);
                            }} else {{
                                element.textContent = value; // For non-input elements
                            }}
                        }});

                        // Dynamic rendering for list sections (experience, education, projects, etc.)
                        // This part assumes the template has a container for each list section, e.g., <div id="experience_list_container"></div>
                        const listSections = {{
                            'experience': {{
                                default: {{ position: '', company: '', duration: '', location: '', achievements: [], technologies: [] }},
                                render: (item, index) => `
                                    <label>Position:</label><input type="text" class="editable-field" data-path="experience[${{index}}].position" value="${{escapeHtml(item.position || '')}}"><br>
                                    <label>Company:</label><input type="text" class="editable-field" data-path="experience[${{index}}].company" value="${{escapeHtml(item.company || '')}}"><br>
                                    <label>Duration:</label><input type="text" class="editable-field" data-path="experience[${{index}}].duration" value="${{escapeHtml(item.duration || '')}}"><br>
                                    <label>Location:</label><input type="text" class="editable-field" data-path="experience[${{index}}].location" value="${{escapeHtml(item.location || '')}}"><br>
                                    <label>Achievements:</label><textarea class="editable-field" data-path="experience[${{index}}].achievements">${{escapeHtml(Array.isArray(item.achievements) ? item.achievements.join('\\n') : item.achievements || '')}}</textarea><br>
                                    <label>Technologies:</label><input type="text" class="editable-field" data-path="experience[${{index}}].technologies" value="${{escapeHtml(Array.isArray(item.technologies) ? item.technologies.join(', ') : item.technologies || '')}}"><br>
                                    <button class="remove-item-button" data-index="${{index}}" data-section="experience">Remove</button>
                                `
                            }},
                            'education': {{
                                default: {{ degree: '', institution: '', graduation_date: '', gpa: '', relevant_coursework: [] }},
                                render: (item, index) => `
                                    <label>Degree:</label><input type="text" class="editable-field" data-path="education[${{index}}].degree" value="${{escapeHtml(item.degree || '')}}"><br>
                                    <label>Institution:</label><input type="text" class="editable-field" data-path="education[${{index}}].institution" value="${{escapeHtml(item.institution || '')}}"><br>
                                    <label>Graduation Date:</label><input type="text" class="editable-field" data-path="education[${{index}}].graduation_date" value="${{escapeHtml(item.graduation_date || '')}}"><br>
                                    <label>GPA:</label><input type="text" class="editable-field" data-path="education[${{index}}].gpa" value="${{escapeHtml(item.gpa || '')}}"><br>
                                    <label>Coursework:</label><textarea class="editable-field" data-path="education[${{index}}].relevant_coursework">${{escapeHtml(Array.isArray(item.relevant_coursework) ? item.relevant_coursework.join('\\n') : item.relevant_coursework || '')}}</textarea><br>
                                    <button class="remove-item-button" data-index="${{index}}" data-section="education">Remove</button>
                                `
                            }},
                            'projects': {{
                                default: {{ name: '', description: '', technologies: [], links: [] }},
                                render: (item, index) => `
                                    <label>Project Name:</label><input type="text" class="editable-field" data-path="projects[${{index}}].name" value="${{escapeHtml(item.name || '')}}"><br>
                                    <label>Description:</label><textarea class="editable-field" data-path="projects[${{index}}].description">${{escapeHtml(item.description || '')}}</textarea><br>
                                    <label>Technologies:</label><input type="text" class="editable-field" data-path="projects[${{index}}].technologies" value="${{escapeHtml(Array.isArray(item.technologies) ? item.technologies.join(', ') : item.technologies || '')}}"><br>
                                    <label>Links:</label><input type="text" class="editable-field" data-path="projects[${{index}}].links" value="${{escapeHtml(Array.isArray(item.links) ? item.links.join(', ') : item.links || '')}}"><br>
                                    <button class="remove-item-button" data-index="${{index}}" data-section="projects">Remove</button>
                                `
                            }},
                            'certifications': {{
                                default: {{ name: '', issuer: '', date: '', expiry: '', link: '' }},
                                render: (item, index) => `
                                    <label>Certification Name:</label><input type="text" class="editable-field" data-path="certifications[${{index}}].name" value="${{escapeHtml(item.name || '')}}"><br>
                                    <label>Issuer:</label><input type="text" class="editable-field" data-path="certifications[${{index}}].issuer" value="${{escapeHtml(item.issuer || '')}}"><br>
                                    <label>Date:</label><input type="text" class="editable-field" data-path="certifications[${{index}}].date" value="${{escapeHtml(item.date || '')}}"><br>
                                    <label>Expiry:</label><input type="text" class="editable-field" data-path="certifications[${{index}}].expiry" value="${{escapeHtml(item.expiry || '')}}"><br>
                                    <label>Link:</label><input type="text" class="editable-field" data-path="certifications[${{index}}].link" value="${{escapeHtml(item.link || '')}}"><br>
                                    <button class="remove-item-button" data-index="${{index}}" data-section="certifications">Remove</button>
                                `
                            }},
                            'achievements_and_awards': {{
                                default: '',
                                render: (item, index) => `
                                    <textarea class="editable-field" data-path="achievements_and_awards[${{index}}]">${{escapeHtml(item || '')}}</textarea><br>
                                    <button class="remove-item-button" data-index="${{index}}" data-section="achievements_and_awards">Remove</button>
                                `
                            }},
                            'languages': {{
                                default: '',
                                render: (item, index) => `
                                    <input type="text" class="editable-field" data-path="languages[${{index}}]" value="${{escapeHtml(item || '')}}"><br>
                                    <button class="remove-item-button" data-index="${{index}}" data-section="languages">Remove</button>
                                `
                            }}
                        }};

                        for (const sectionKey in listSections) {{
                            const container = document.getElementById(`${{sectionKey}}_list_container`);
                            if (container && window.resumeData[sectionKey]) {{
                                container.innerHTML = ''; // Clear existing content
                                window.resumeData[sectionKey].forEach((item, index) => {{
                                    const itemDiv = document.createElement('div');
                                    itemDiv.className = 'editable-list-item';
                                    itemDiv.innerHTML = listSections[sectionKey].render(item, index);
                                    container.appendChild(itemDiv);
                                }});
                                const addButton = document.createElement('button');
                                addButton.className = 'add-item-button';
                                addButton.textContent = `Add ${{sectionKey.replace(/_/g, ' ').replace(/\\b[A-Za-z0-9_]/g, l => l.toUpperCase())}}`;
                                addButton.onclick = () => addListItem(sectionKey, listSections[sectionKey].default);
                                container.appendChild(addButton);
                            }}
                        }}

                        // Add event listeners for input changes to update window.resumeData
                        document.querySelectorAll('.editable-field').forEach(input => {{
                            input.addEventListener('change', (event) => {{
                                const dataPath = event.target.dataset.path;
                                let value = event.target.value;

                                // Convert back to array if it was originally an array (e.g., achievements, technologies)
                                if (dataPath.includes('achievements') || dataPath.includes('technologies') || dataPath.includes('links') || dataPath.includes('relevant_coursework')) {{
                                    value = value.split('\\n').map(s => s.trim()).filter(s => s.length > 0);
                                }} else if (dataPath.includes('skills')) {{
                                    // For skills, split by comma and trim
                                    value = value.split(',').map(s => s.trim()).filter(s => s.length > 0);
                                }}

                                setObjectValue(window.resumeData, dataPath, value);
                                console.log('Updated resumeData:', window.resumeData);
                            }});
                        }});

                        // Re-attach remove button listeners after re-rendering
                        document.querySelectorAll('.remove-item-button').forEach(button => {{
                            button.onclick = (event) => {{
                                const index = parseInt(event.target.dataset.index);
                                const section = event.target.dataset.section;
                                removeListItem(section, index);
                            }};
                        }});
                    }}

                    function addListItem(section, defaultValues) {{
                        if (!window.resumeData[section]) {{
                            window.resumeData[section] = [];
                        }}
                        window.resumeData[section].push(defaultValues);
                        renderEditableResume(); // Re-render to show new item
                    }}

                    function removeListItem(section, index) {{
                        if (confirm('Are you sure you want to remove this item?')) {{
                            window.resumeData[section].splice(index, 1);
                            renderEditableResume(); // Re-render to reflect removal
                        }}
                    }}

                    function escapeHtml(text) {{
                        const map = {{
                            '&': '&',
                            '<': '<',
                            '>': '>',
                            '"': '"',
                            "'": '&#039;',
                            '`': '&#x60;'
                        }};
                        return text.replace(/[&<>"'`]/g, function(m) {{ return map[m]; }});
                    }}

                    document.addEventListener('DOMContentLoaded', renderEditableResume);

                    {escaped_js_content}
                </script>
            </body>
            </html>
        """

        # Save the content to a temporary HTML file in the 'uploads' directory
        temp_filename = f"preview_{uuid.uuid4()}.html"
        temp_filepath = os.path.join(current_app.config['UPLOAD_FOLDER'], temp_filename)

        with open(temp_filepath, "w", encoding="utf-8") as f:
            f.write(iframe_html_content)

        # Return the URL to the temporary file
        preview_url = f"/uploads/{temp_filename}"
        return jsonify({"preview_url": preview_url, "parsed_data": parsed_data, "parsed_data_id": parsed_data_id}), 200

    except Exception as e:
        logger.error(f"Error generating resume preview: {e}")
        return jsonify({"error": "Internal server error while generating preview."}), 500
    finally:
        client.close()

@app.route('/uploads/<filename>')
def uploaded_file(filename):
    """Serve uploaded files from the UPLOAD_FOLDER"""
    return send_from_directory(app.config['UPLOAD_FOLDER'], filename)

@app.route('/health')
def health_check():
    """Health check endpoint"""
    return jsonify({
        'status': 'healthy',
        'timestamp': datetime.now().isoformat(),
        'mistral_ocr_available': mistral_client is not None,
        'mistral_api_configured': MISTRAL_API_KEY is not None,
        'spacy_loaded': nlp is not None
    }), 200

if __name__ == '__main__':
    app.run(debug=True, host='127.0.0.1', port=5000)